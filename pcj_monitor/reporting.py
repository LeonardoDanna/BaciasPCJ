from __future__ import annotations

import csv
import json
import re
from collections import defaultdict
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from typing import Iterable

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from docx.shared import Pt
except ImportError:
    Pt = None

from .analyzer import build_daily_summary, sort_key
from .config import MonitorConfig
from .models import RelevantNews


def add_bold_label(paragraph, label: str, value: str) -> None:
    run = paragraph.add_run(label)
    run.bold = True
    paragraph.add_run(value)


def add_highlighted_text(paragraph, text: str, terms: Iterable[str]) -> None:
    filtered_terms = sorted({term.strip() for term in terms if term and term.strip()}, key=len, reverse=True)
    if not filtered_terms:
        paragraph.add_run(text)
        return

    pattern = re.compile("(" + "|".join(re.escape(term) for term in filtered_terms) + ")", re.IGNORECASE)
    last = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last:
            paragraph.add_run(text[last:start])
        run = paragraph.add_run(text[start:end])
        run.bold = True
        last = end
    if last < len(text):
        paragraph.add_run(text[last:])


def generate_docx_report(
    items: list[RelevantNews],
    generated_at: datetime,
    output_path: Path,
    config: MonitorConfig,
    source_stats: list[dict[str, object]],
    topic_summary: dict | None = None,
) -> None:
    if Document is None:
        raise RuntimeError(
            "A biblioteca 'python-docx' nao esta instalada. Instale com: pip3 install python-docx"
        )

    document = Document()
    if Pt is not None:
        styles = document.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)
        styles["Heading 1"].font.name = "Arial"
        styles["Heading 2"].font.name = "Arial"
        styles["Title"].font.name = "Arial"

    document.add_heading("Relatorio Diario de Noticias Hidricas - Bacias PCJ", level=0)
    meta = document.add_paragraph()
    add_bold_label(meta, "Gerado em: ", generated_at.strftime("%d/%m/%Y %H:%M:%S UTC"))

    document.add_heading("Resumo Consolidado", level=1)
    for line in build_daily_summary(items, config).splitlines():
        paragraph = document.add_paragraph(style="List Bullet")
        add_highlighted_text(paragraph, line, {"Alta", "Media", "Baixa"} | config.keywords_b | config.keywords_a)

    if topic_summary and topic_summary.get('total_topics', 0) > 0:
        document.add_heading("Análise de Tópicos (BERTopic)", level=1)
        document.add_paragraph(f"Total de tópicos identificados: {topic_summary['total_topics']}")
        for topic in topic_summary['topics']:
            document.add_heading(f"Tópico {topic['topic_id']}: {topic['name']}", level=2)
            p = document.add_paragraph()
            add_bold_label(p, "Contagem: ", str(topic['count']))
            p = document.add_paragraph()
            add_bold_label(p, "Palavras principais: ", ", ".join(topic['representation'][:10]))

    document.add_heading("Execucao", level=1)
    success_count = sum(1 for stat in source_stats if str(stat.get("status", "")).startswith(("feed", "portal", "artigo")))
    exec_paragraph = document.add_paragraph(style="List Bullet")
    add_bold_label(exec_paragraph, "Fontes processadas: ", f"{success_count}/{len(source_stats)}")
    for stat in source_stats:
        paragraph = document.add_paragraph(style="List Bullet")
        add_bold_label(
            paragraph,
            f"{stat['url']}: ",
            f"status={stat.get('status')} | noticias={stat.get('articles_found', 0)}",
        )

    grouped: dict[str, list[RelevantNews]] = defaultdict(list)
    for item in sorted(items, key=lambda news: sort_key(news, config)):
        grouped[item.municipio or "Municipio nao identificado"].append(item)

    for municipality, news in grouped.items():
        document.add_heading(municipality, level=1)
        for item in news:
            document.add_heading(item.titulo, level=2)
            p = document.add_paragraph()
            add_bold_label(p, "Data: ", item.data)
            p = document.add_paragraph()
            add_bold_label(p, "Municipio: ", item.municipio or "Nao identificado")
            p = document.add_paragraph()
            add_bold_label(p, "Fonte: ", item.fonte)
            p = document.add_paragraph()
            add_bold_label(p, "Link: ", item.link)
            p = document.add_paragraph()
            add_bold_label(p, "Criticidade: ", item.classificacao)
            p = document.add_paragraph()
            add_bold_label(p, "Palavras-chave detectadas: ", ", ".join(item.palavras_chave_detectadas))
            p = document.add_paragraph()
            label = p.add_run("Resumo: ")
            label.bold = True
            for line in item.resumo.splitlines():
                paragraph = document.add_paragraph(style="List Bullet")
                add_highlighted_text(paragraph, line, item.palavras_chave_detectadas)
            document.add_paragraph("")

    document.save(output_path)


def build_json_payload(
    items: list[RelevantNews],
    generated_at: datetime,
    config: MonitorConfig,
    source_stats: list[dict[str, object]],
    topic_summary: dict | None = None,
) -> dict[str, object]:
    return {
        "generated_at": generated_at.isoformat(),
        "summary": build_daily_summary(items, config),
        "total_relevant_news": len(items),
        "source_stats": source_stats,
        "news": [asdict(item) for item in items],
        "topic_summary": topic_summary or {},
    }


def write_json_report(output_path: Path, payload: dict[str, object]) -> None:
    output_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def write_csv_report(output_path: Path, payload: dict[str, object]) -> None:
    fieldnames = [
        "data",
        "municipio",
        "fonte",
        "titulo",
        "link",
        "resumo",
        "classificacao",
        "palavras_chave_detectadas",
        "entidades_locais",
        "eventos_problemas",
        "topico",
    ]

    with output_path.open("w", encoding="utf-8", newline="") as csv_file:
        writer = csv.DictWriter(csv_file, fieldnames=fieldnames)
        writer.writeheader()
        for item in payload.get("news", []):
            writer.writerow(
                {
                    "data": item.get("data", ""),
                    "municipio": item.get("municipio", "") or "",
                    "fonte": item.get("fonte", ""),
                    "titulo": item.get("titulo", ""),
                    "link": item.get("link", ""),
                    "resumo": item.get("resumo", ""),
                    "classificacao": item.get("classificacao", ""),
                    "palavras_chave_detectadas": ", ".join(item.get("palavras_chave_detectadas", [])),
                    "entidades_locais": ", ".join(item.get("entidades_locais", [])),
                    "eventos_problemas": ", ".join(item.get("eventos_problemas", [])),
                    "topico": item.get("topico", -1),
                }
            )
